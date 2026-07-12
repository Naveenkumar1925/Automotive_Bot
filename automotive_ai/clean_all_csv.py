import os
import csv
import pandas as pd
from docx import Document

# Windows Word automation
import win32com.client


# =========================
# SETTINGS
# =========================

input_folder = r"data/raw"
output_folder = r"data/output_text"

os.makedirs(output_folder, exist_ok=True)


# =========================
# HELPERS
# =========================

def write_with_spacing(lines, output_path):

    with open(output_path, "w", encoding="utf-8") as f:
        for line in lines:
            f.write(line + "\n\n")   # empty line between rows


def detect_encoding(path):

    encodings = ["utf-8-sig", "cp1252", "latin1"]

    for enc in encodings:
        try:
            with open(path, encoding=enc) as f:
                f.read()
            return enc
        except:
            continue

    return "latin1"


def convert_doc_to_docx(doc_path):

    print("Converting .doc → .docx using MS Word...")

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    doc = word.Documents.Open(os.path.abspath(doc_path))

    new_path = doc_path + "x"   # .doc → .docx
    doc.SaveAs(os.path.abspath(new_path), 16)

    doc.Close()
    word.Quit()

    return new_path


# =========================
# MAIN LOOP
# =========================

for filename in os.listdir(input_folder):

    input_path = os.path.join(input_folder, filename)

    if not os.path.isfile(input_path):
        continue

    name, ext = os.path.splitext(filename)
    ext = ext.lower()

    output_path = os.path.join(output_folder, name + ".txt")

    print(f"Processing: {filename}")

    try:

        # =====================
        # CSV
        # =====================
        if ext == ".csv":

            encoding = detect_encoding(input_path)
            lines = []

            with open(input_path, newline='', encoding=encoding) as f:
                reader = csv.reader(f)
                for row in reader:
                    lines.append(",".join(row))

            write_with_spacing(lines, output_path)


        # =====================
        # XLSX / XLS
        # =====================
        elif ext in [".xlsx", ".xls"]:

            df = pd.read_excel(input_path)

            if df.empty:
                print("⚠ Empty Excel file")
                continue

            df = df.fillna("")
            lines = df.astype(str).agg(",".join, axis=1).tolist()

            write_with_spacing(lines, output_path)


    
        elif ext == ".docx":

            doc = Document(input_path)

            lines = []

            for p in doc.paragraphs:
                text = p.text.strip()
                if text:
                    lines.append(text)

            write_with_spacing(lines, output_path)


        # =====================
        # DOC (auto convert)
        # =====================
        elif ext == ".doc":

            new_docx = convert_doc_to_docx(input_path)

            doc = Document(new_docx)

            lines = []

            for p in doc.paragraphs:
                text = p.text.strip()
                if text:
                    lines.append(text)

            write_with_spacing(lines, output_path)


        else:
            print("Skipped unsupported file:", filename)

    except Exception as e:
        print(f"❌ Failed {filename}: {e}")


print("✅ ALL FILES CONVERTED SUCCESSFULLY")